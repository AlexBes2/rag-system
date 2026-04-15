from __future__ import annotations

import os
from pathlib import Path
from typing import Any

from docx import Document
from pdf2image import convert_from_path
from PIL import Image, ImageFilter, ImageOps, UnidentifiedImageError
from pypdf import PdfReader
import pytesseract

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".png",
    ".jpg",
    ".jpeg",
    ".tif",
    ".tiff",
    ".bmp",
}

IMAGE_EXTENSIONS = {
    ".png",
    ".jpg",
    ".jpeg",
    ".tif",
    ".tiff",
    ".bmp",
}

OCR_LANGS = os.getenv("OCR_LANGS", "ukr+rus+eng")
OCR_DPI = int(os.getenv("OCR_DPI", "300"))
OCR_MIN_TEXT_LEN = int(os.getenv("OCR_MIN_TEXT_LEN", "30"))
OCR_TIMEOUT = float(os.getenv("OCR_TIMEOUT", "60"))
TESSERACT_CMD = os.getenv("TESSERACT_CMD", "").strip()

if TESSERACT_CMD:
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD


def _normalize_text(text: str) -> str:
    return " ".join((text or "").split()).strip()


def _needs_ocr(text: str) -> bool:
    return len(_normalize_text(text)) < OCR_MIN_TEXT_LEN


def _preprocess_image(image: Image.Image) -> Image.Image:
    image = ImageOps.exif_transpose(image)

    if image.mode != "RGB":
        image = image.convert("RGB")

    gray = ImageOps.grayscale(image)
    gray = ImageOps.autocontrast(gray)

    width, height = gray.size
    if max(width, height) < 1800:
        gray = gray.resize((width * 2, height * 2), Image.LANCZOS)

    gray = gray.filter(ImageFilter.SHARPEN)

    bw = gray.point(lambda x: 0 if x < 160 else 255, mode="1")
    return bw.convert("RGB")


def _ocr_image(image: Image.Image) -> str:
    processed = _preprocess_image(image)

    try:
        text = pytesseract.image_to_string(
            processed,
            lang=OCR_LANGS,
            config="--oem 1 --psm 6",
            timeout=OCR_TIMEOUT,
        )
    except RuntimeError as exc:
        raise ValueError(f"OCR завершился с ошибкой или по таймауту: {exc}") from exc

    return text.strip()


def _build_section(
    *,
    document_name: str,
    page: int,
    section_index: int,
    text: str,
    source_type: str,
    ocr_used: bool,
) -> dict[str, Any]:
    return {
        "document_name": document_name,
        "page": page,
        "section_index": section_index,
        "text": text.strip(),
        "source_type": source_type,
        "ocr_used": ocr_used,
    }


def _extract_pdf_sections(path: Path) -> list[dict[str, Any]]:
    reader = PdfReader(str(path))
    document_name = path.name

    sections: list[dict[str, Any]] = []
    pages_needing_ocr: list[int] = []

    for page_number, page in enumerate(reader.pages, start=1):
        # Сначала пробуем layout extraction — для колонок обычно лучше
        try:
            extracted_text = (page.extract_text(extraction_mode="layout") or "").strip()
        except TypeError:
            # На случай старой версии pypdf
            extracted_text = (page.extract_text() or "").strip()

        # Если layout extraction дал совсем мало текста, пробуем обычный режим
        if _needs_ocr(extracted_text):
            fallback_text = (page.extract_text() or "").strip()
            if len(_normalize_text(fallback_text)) > len(_normalize_text(extracted_text)):
                extracted_text = fallback_text

        sections.append(
            _build_section(
                document_name=document_name,
                page=page_number,
                section_index=page_number - 1,
                text=extracted_text,
                source_type="pdf",
                ocr_used=False,
            )
        )

        if _needs_ocr(extracted_text):
            pages_needing_ocr.append(page_number)

    if pages_needing_ocr:
        images = convert_from_path(str(path), dpi=OCR_DPI)

        for page_number in pages_needing_ocr:
            image = images[page_number - 1]
            ocr_text = _ocr_image(image)
            current_text = sections[page_number - 1]["text"]

            if len(_normalize_text(ocr_text)) >= len(_normalize_text(current_text)):
                sections[page_number - 1]["text"] = ocr_text

            sections[page_number - 1]["ocr_used"] = bool(_normalize_text(ocr_text))

    return [section for section in sections if section["text"].strip()]


def _extract_docx_sections(path: Path) -> list[dict[str, Any]]:
    document = Document(str(path))
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = " ".join(part.strip() for part in cell.text.splitlines() if part.strip())
                if cell_text:
                    row_cells.append(cell_text)

            if row_cells:
                lines.append(" | ".join(row_cells))

    full_text = "\n".join(lines).strip()
    if not full_text:
        return []

    return [
        _build_section(
            document_name=path.name,
            page=1,
            section_index=0,
            text=full_text,
            source_type="docx",
            ocr_used=False,
        )
    ]


def _extract_image_sections(path: Path) -> list[dict[str, Any]]:
    try:
        image = Image.open(path)
    except UnidentifiedImageError as exc:
        raise ValueError(f"Не удалось открыть изображение: {path.name}") from exc

    text = _ocr_image(image)
    if not text.strip():
        return []

    return [
        _build_section(
            document_name=path.name,
            page=1,
            section_index=0,
            text=text,
            source_type="image",
            ocr_used=True,
        )
    ]


def extract_document_sections(path: Path) -> list[dict[str, Any]]:
    if not path.exists():
        raise ValueError(f"Файл не найден: {path}")

    extension = path.suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(f"Поддерживаются только: {supported}")

    if extension == ".pdf":
        return _extract_pdf_sections(path)

    if extension == ".docx":
        return _extract_docx_sections(path)

    if extension in IMAGE_EXTENSIONS:
        return _extract_image_sections(path)

    return []